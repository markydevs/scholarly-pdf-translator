#!/usr/bin/env python3

from __future__ import annotations
import argparse, sys, re, time, random, tempfile, shutil, subprocess
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any

# ---------- extraction ----------
try:
    from pdfminer.high_level import extract_text  # type: ignore
    _PDFMINER = True
except Exception:
    _PDFMINER = False
    extract_text = None  # type: ignore

def _extract_all_text(pdf: Path) -> str:
    if _PDFMINER and extract_text is not None:
        return (extract_text(str(pdf)) or "").strip()
    try:
        import PyPDF2  # type: ignore
    except Exception:
        raise RuntimeError("Install pdfminer.six or PyPDF2 to extract text.")
    out = []
    for p in PyPDF2.PdfReader(str(pdf)).pages:
        try:
            out.append(p.extract_text() or "")
        except Exception:
            pass
    return "\n\f\n".join(out).strip()

# ---------- tiny utils ----------
def _run(cmd: List[str], passthrough: bool = False) -> None:
    if passthrough:
        subprocess.run(cmd, check=True)
    else:
        p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if p.returncode != 0:
            raise RuntimeError(f"Command failed: {' '.join(cmd)}\nSTDERR:\n{p.stderr}")

def safe_stem(path: Path) -> str:
    return re.sub(r"[^\w\-]+", "_", path.stem)

def paragraphs_from_text(txt: str) -> List[str]:
    txt = re.sub(r'[ \t]*\n[ \t]*', ' ', txt)
    parts = re.split(r'\n\s*\n+', txt)
    return [p.strip() for p in parts if p.strip()]

def chunk_text(paragraphs: List[str], max_chars: int) -> List[str]:
    chunks, buf, total = [], [], 0
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if len(p) > max_chars:
            cur = ""
            for s in re.split(r'(?<=[.!?;:])\s+', p):
                if len(cur) + len(s) + 1 > max_chars:
                    if cur: chunks.append(cur)
                    cur = s
                else:
                    cur = (cur + " " + s) if cur else s
            if cur: chunks.append(cur)
            continue
        if total + len(p) + 1 > max_chars and buf:
            chunks.append("\n\n".join(buf))
            buf, total = [p], len(p)
        else:
            buf.append(p); total += len(p) + 1
    if buf: chunks.append("\n\n".join(buf))
    return chunks

# ---------- cleanup / structure ----------
LIGS = {"ﬁ":"fi","ﬂ":"fl","ﬃ":"ffi","ﬄ":"ffl","ﬅ":"ft","ﬆ":"st"}
PAGE_NUM_RE = re.compile(r"^\s*([ivxlcdm]+|\d+)\s*$", re.I)
ALLCAPS_SHORT_RE = re.compile(r"^[^a-z]{1,48}$")
HEADER_FOOTER_HINTS = [
    "CORPUS","SCRIPTORUM","CHRISTIANORUM","ORIENTALIUM",
    "SCRIPTORES SYRI","IMPRIMERIE","LOUVAIN","RÉIMPRESSION","REIMPRESSION","ANASTATIQUE",
    "UNIVERSITATIS","CATHOLICAE","AMERICAE","LOVANIENSIS",
    "ORATIO","LIBER","GRAMMATICUM","ORATIONIS","PARS POSTERIOR",
    "ORIGINAL","TRANSLATION","FOOTNOTES",  # table column labels
    " P:"," A "," B "," C "               # side/column letters
]

# Apparatus tokens (robust)
_APP_SIGLA = r"(?:P\.\s*Gr\.?|PG|PL|CSCO|CSEL|PO|BHG|BHL|CPG|Ibid\.?|Id\.?|Cf\.|Cfr\.)"
APPARATUS_LINE   = re.compile(rf"^\s*(?:[—–\-*·•;:,]+)?\s*(?:{_APP_SIGLA}).*$", re.I)
APPARATUS_INLINE = re.compile(rf"(?:(?:{_APP_SIGLA})[^\n.;]*[.;]?)", re.I)

# Build a strict note-start regex  (p. N is optional, off by default)
def build_note_start_re(include_p_cue: bool = False) -> re.Pattern[str]:
    p_cue = r"(?:\*+\s*)?p\.\s*\d{1,3}\.?"
    core = rf"""^\s*
        (?:[;:,·•]*)?
        (?:
             \(?\d{{1,3}}\)?[.)]     # (1) / 1) / 1.
           | [a-z]\)                 # a)
           | [†‡]+                   # daggers
           | (?:{_APP_SIGLA})\b      # PG/PL/CSCO/CSEL/.../Ibid./Cf.
           {"| " + p_cue if include_p_cue else ""}
        )
        \s+
    """
    return re.compile(core, re.I | re.X | re.U)

# Body-line junk like "— * — ^ —" at the beginning
NOISE_PREFIX_RE = re.compile(r"^\s*[—–\-*^_·•\"'`~]+(?:\s*[—–\-*^_·•\"'`~]+)*\s*(?=[A-Za-z\[\(0-9\u0370-\u03FF\u1F00-\u1FFF])")

_LATIN_ENDINGS = {"M","UM","I","IS","OS","US","AE","A","AS","ES","EM","AM","ORUM","ARUM","U","IUM","IBUS"}
_caps_end_pat = re.compile(r"\b([A-Z]{3,})\s+([A-Z]{1,4})\b")

def _fix_caps_line(line: str) -> str:
    s = line.strip()
    if not s or s != s.upper():
        return line
    cur = line
    for _ in range(3):
        prev = cur
        cur = _caps_end_pat.sub(
            lambda m: (m.group(1)+m.group(2)) if (m.group(2) in _LATIN_ENDINGS or len(m.group(2))==1) else m.group(0),
            cur
        )
        if cur == prev: break
    return cur

def _looks_front_matter(text: str) -> bool:
    letters = re.findall(r"[A-Za-z]", text)
    if not letters: return True
    lower_ratio = sum(ch.islower() for ch in letters) / max(1,len(letters))
    up = text.upper()
    return lower_ratio < 0.06 and any(h in up for h in HEADER_FOOTER_HINTS)

# ---- apparatus normalization helpers ----
_ROMAN_PAIR = re.compile(r"\b([IVXLCDM]{1,6})\s+([IVXLCDM]{1,6})\b")
def _compact_roman_numerals(s: str) -> str:
    return _ROMAN_PAIR.sub(lambda m: m.group(1)+m.group(2), s)

def _norm_apparatus_text(s: str) -> str:
    s = s.strip()
    s = re.sub(r"^[—–\-*·•;:,]+\s*", "", s)
    s = re.sub(r"\s+", " ", s)
    s = _compact_roman_numerals(s)
    return s.strip(" ;,·")

def _strip_editorial_and_headers(t: str, extra_hints: Optional[List[str]], apparatus_mode: str) -> Tuple[str, Optional[str]]:
    """
    Remove running heads/footers & collect apparatus.
    We DO NOT globally siphon generic note-looking lines here.
    """
    foot_accum: List[str] = []
    hints = [*HEADER_FOOTER_HINTS, *(h.upper() for h in (extra_hints or []))]

    lines = t.splitlines()
    lines_out: List[str] = []

    in_app_block = False
    app_buf: List[str] = []

    def flush_app():
        nonlocal app_buf
        if app_buf:
            joined = " ".join(_norm_apparatus_text(x) for x in app_buf if _norm_apparatus_text(x))
            if joined:
                foot_accum.append(joined)
        app_buf = []

    for raw in lines:
        ln = raw.rstrip()
        if not ln:
            if in_app_block:
                flush_app()
                in_app_block = False
            continue

        if PAGE_NUM_RE.match(ln):
            continue

        up = ln.upper()
        # running heads/footers & short all-caps junk
        if (ALLCAPS_SHORT_RE.match(ln) and any(h in up for h in hints)) or (any(h in up for h in hints) and len(ln) < 90):
            continue

        # SIPHON ONLY TRUE APPARATUS LINES ANYWHERE
        if apparatus_mode in {"drop","footnotes"} and APPARATUS_LINE.match(ln):
            if apparatus_mode == "footnotes":
                in_app_block = True
                app_buf.append(ln)
            continue

        # Inline apparatus phrases → footnotes / removed
        if apparatus_mode in {"drop","footnotes"} and APPARATUS_INLINE.search(ln):
            if apparatus_mode == "footnotes":
                for m in APPARATUS_INLINE.finditer(ln):
                    cleaned = _norm_apparatus_text(m.group(0))
                    if cleaned:
                        foot_accum.append(cleaned)
            ln = APPARATUS_INLINE.sub("", ln).strip()

        if in_app_block:
            flush_app()
            in_app_block = False

        # strip noisy punctuation runs at start of body lines
        ln = NOISE_PREFIX_RE.sub("", ln)
        lines_out.append(_fix_caps_line(ln))

    if in_app_block:
        flush_app()

    cleaned = "\n".join(lines_out)
    foot = ("\n".join([x for x in foot_accum if x]).strip() or None)
    return cleaned, foot

def split_body_and_footnotes(
    page_text: str,
    zone_ratio: float,
    min_lines: int,
    density_req: float,
    note_start_re: re.Pattern[str],
    mode: str = "strict",
) -> Tuple[str, Optional[str]]:
    """
    mode: "strict" (default) or "none"
    - "none": never split; all text stays in body (apparatus siphon still applies).
    - "strict": bottom-window split using conservative heuristics.
    """
    if mode == "none":
        lines = [ln.rstrip() for ln in page_text.splitlines() if ln.strip()]
        return ("\n".join(lines).strip(), None)

    lines = [ln.rstrip() for ln in page_text.splitlines() if ln.strip()]
    if not lines:
        return "", None

    n = len(lines)
    start = max(0, int(n * (1 - max(0.0, min(zone_ratio, 1.0)))))
    window = lines[start:] or []

    def is_note(ln: str) -> bool:
        # Long/flowing lines are typically body, not note heads
        if len(ln) > 170:
            return False
        return bool(note_start_re.match(ln) or APPARATUS_LINE.match(ln))

    note_flags = [is_note(x) for x in window]
    density = (sum(note_flags) / max(1, len(note_flags)))
    if density < density_req:
        return "\n".join(lines).strip(), None

    # find first contiguous run of notes from bottom
    i = n - 1
    run_start = None
    run_len = 0
    while i >= start:
        if is_note(lines[i]):
            run_start = i
            run_len += 1
            i -= 1
            while i >= start and is_note(lines[i]):
                run_len += 1
                run_start = i
                i -= 1
            break
        i -= 1

    if run_start is not None and run_len >= max(2, min_lines):
        body = "\n".join(lines[:run_start]).strip()
        foot = "\n".join(lines[run_start:]).strip() or None
        return body, foot

    return "\n".join(lines).strip(), None

_CAPUT_RE = re.compile(r"^\s*CAPUT\s+([A-Z\s\-]+)\s*$")
LATIN_ORD = {
    "PRIMUM":1,"SECUNDUM":2,"TERTIUM":3,"QUARTUM":4,"QUINTUM":5,"SEXTUM":6,"SEPTIMUM":7,"OCTAVUM":8,"NONUM":9,"DECIMUM":10,
    "UNDECIMUM":11,"DUODECIMUM":12,"TERTIUM DECIMUM":13,"QUARTUM DECIMUM":14,"QUINTUM DECIMUM":15,"SEX-DECIMUM":16,"SEPTIMUM DECIMUM":17,
    "DUODEVICESIMUM":18,"UNDEVICESIMUM":19,"VICESIMUM":20,"VICESIMUM PRIMUM":21,"VICESIMUM ALTERUM":22,"VICESIMUM TERTIUM":23,
    "VICESIMUM QUARTUM":24,"VICESIMUM QUINTUM":25,"VICESIMUM SEXTUM":26,"VICESIMUM SEPTIMUM":27,"VICESIMUM OCTAVUM":28,"VICESIMUM NONUM":29,
    "TRICESIMUM":30
}
def _maybe_normalize_caput(line: str, normalize: bool) -> Optional[str]:
    if not normalize: return None
    m = _CAPUT_RE.match(line.upper())
    if not m: return None
    key = " ".join(m.group(1).split())
    n = LATIN_ORD.get(key)
    return f"CHAPTER {n}" if n is not None else None

# --- marginal line-number stripping ---
MARGIN_NUM_RE = re.compile(r"^\s*(?:\d{1,3})(?:\s+\d{1,3}){0,6}[.:;]?\s*$")
def _strip_marginal_numbers(text: str) -> str:
    out = []
    for ln in text.splitlines():
        s = ln.strip()
        if MARGIN_NUM_RE.match(s):
            continue
        out.append(ln)
    return "\n".join(out)

# --- heading detection ---
def detect_title(page_text: str) -> Optional[str]:
    def _upper_ratio(s: str) -> float:
        letters = [c for c in s if c.isalpha()]
        if not letters: return 0.0
        return sum(c.isupper() for c in letters) / len(letters)

    lines = [ln.strip(" |·—–").strip() for ln in page_text.splitlines() if ln.strip()]
    head_lines = lines[:8]

    for ln in head_lines:
        if PAGE_NUM_RE.match(ln): continue
        if ln.lower() in {"original","translation","footnotes"}: continue
        up = ln.upper()
        is_all_capsish = _upper_ratio(ln) >= 0.85 and len(ln) <= 80
        is_caput = up.startswith("CAPUT") and len(ln) <= 120
        if (is_all_capsish or is_caput) and not re.search(r"[.!?]$", ln):
            return ln

    join3 = " ".join(head_lines[:3])
    m = re.search(r"([A-Z][A-Z0-9 .,:;—–-]{4,300}?CAPUT\s+[A-Z \-]+)", join3)
    if m:
        return m.group(1).strip()
    return None

def peel_heading_block(text: str) -> Tuple[Optional[str], str]:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines: return None, text

    def ok_head(s: str) -> bool:
        s = s.strip(" |·—–").strip()
        if not s: return False
        letters = [c for c in s if s and c.isalpha()]
        upper_ratio = (sum(c.isupper() for c in letters) / len(letters)) if letters else 0.0
        return (s.upper().startswith("CAPUT") and len(s) <= 120) or (upper_ratio >= 0.85 and len(s) <= 90)

    head, i = [], 0
    for ln in lines[:6]:
        if ok_head(ln):
            head.append(ln.strip("| ").strip()); i += 1
            if i == 3: break
        else:
            break
    if head:
        return " ".join(head), "\n".join(lines[i:])

    join3 = " ".join(lines[:3]).strip(" |·—–").strip()
    m = re.match(r"(?P<head>[A-Z][A-Z0-9 .,:;—–-]{4,300}?CAPUT\s+[A-Z \-]+)\s+(?P<body>[A-Z].*)$", join3)
    if m:
        head_text = m.group("head").strip()
        drop, acc = 0, ""
        for ln in lines[:3]:
            acc = (acc + " " + ln.strip()).strip()
            drop += 1
            if head_text in acc:
                break
        body = m.group("body") + ("\n" + "\n".join(lines[drop:]) if drop < len(lines) else "")
        return head_text, body
    return None, text

# ---------- multi-script masking ----------
SCRIPT_RANGES = {
    "greek": ((0x0370,0x03FF),(0x1F00,0x1FFF)),
    "syriac": ((0x0700,0x074F),),
    "hebrew": ((0x0590,0x05FF),),
    "coptic": ((0x2C80,0x2CFF),),
    "armenian": ((0x0530,0x058F),),
}
def _compile_mask_regex(scripts: List[str]) -> re.Pattern[str]:
    ranges = []
    for s in scripts:
        ranges.extend(SCRIPT_RANGES.get(s, ()))
    if not ranges:
        return re.compile("$^")
    parts = "".join([f"\\u{a:04X}-\\u{b:04X}" for (a,b) in ranges])
    return re.compile(f"(?:[{parts}]{{3,}})")  # mask only >=3-length runs

def _mask(text: str, rx: re.Pattern[str]) -> Tuple[str, Dict[str, str]]:
    mapping: Dict[str, str] = {}
    idx = 0
    def repl(m: re.Match[str]) -> str:
        nonlocal idx
        key = f"[[G{idx}]]"; idx += 1
        mapping[key] = m.group(0)
        return key
    return rx.sub(repl, text), mapping

def _unmask(text: str, mapping: Dict[str, str]) -> str:
    for k,v in mapping.items(): text = text.replace(k, v)
    return text

# --- detect real script runs & strip tiny noise -------
def _max_run_len(s: str, ranges: Tuple[Tuple[int,int], ...]) -> int:
    cur = mx = 0
    for ch in s:
        o = ord(ch)
        if any(a <= o <= b for a,b in ranges):
            cur += 1; mx = max(mx, cur)
        else:
            cur = 0
    return mx

def _has_real_greek(s: str) -> bool:
    return _max_run_len(s, tuple(SCRIPT_RANGES["greek"])) >= 3

def _has_real_syriac(s: str) -> bool:
    return _max_run_len(s, tuple(SCRIPT_RANGES["syriac"])) >= 3

def _strip_false_nonlatin_noise(s: str) -> str:
    if not (_has_real_greek(s) or _has_real_syriac(s)):
        s = re.sub(r"[\u0370-\u03FF\u1F00-\u1FFF]{1,2}", "", s)
        s = re.sub(r"[\u0700-\u074F]{1,2}", "", s)
    return s

# ---------- OCR (optional) ----------
def ensure_searchable_pdf(src_pdf: Path, ocr_langs: str, tmpdir: Path,
                          force_all: bool=False, oversample: int=600,
                          psm: int=4, verbose: bool=False) -> Tuple[Path, Optional[Path]]:
    out_pdf = tmpdir / (safe_stem(src_pdf) + "_ocr.pdf")
    sidecar = tmpdir / (safe_stem(src_pdf) + "_ocr.txt")
    cmd = [
        "ocrmypdf","--rotate-pages","--deskew",
        "-l", ocr_langs.replace(" ", "").replace(",", "+"),
        "--oversample", str(oversample),
        "--tesseract-pagesegmode", str(psm),
        "--tesseract-oem", "1",
        "--tesseract-thresholding", "sauvola",
        "--output-type", "pdf",
        "--sidecar", str(sidecar),
        "--force-ocr" if force_all else "--skip-text",
        str(src_pdf), str(out_pdf),
    ]
    if verbose: print("OCR cmd:", " ".join(cmd), flush=True)
    _run(cmd, passthrough=verbose)
    return out_pdf, (sidecar if sidecar.exists() else None)

# --- Smart-OCR heuristic ---
MENTIONS_NONLATIN = re.compile(r"\bP\.\s*Gr\.|\bgraec[ae]\b|\bGreek\b|SYR\.|ܐ|ܡ", re.I)
def _has_nonlatin_codes(s: str) -> bool:
    return _has_real_greek(s) or _has_real_syriac(s)

def _needs_force_ocr(text: str) -> bool:
    t = _strip_false_nonlatin_noise(text)
    mentions = bool(MENTIONS_NONLATIN.search(t))
    has_real = _has_nonlatin_codes(t)
    return mentions and not has_real

# ---------- OpenAI translator ----------
class OpenAITranslator:
    def __init__(self, model: str = "gpt-5-mini", verbose: bool = False,
                 mask_regex: Optional[re.Pattern[str]] = None) -> None:
        try:
            from openai import OpenAI  # type: ignore
        except Exception as e:
            raise RuntimeError("pip install openai") from e
        self.client = OpenAI()
        self.model = model
        self.verbose = verbose
        self.mask_regex = mask_regex or re.compile("$^")

    def translate(self, texts: List[str], src: Optional[str], tgt: str) -> List[str]:
        out: List[str] = []
        sys_msg = (
            "You are an expert translator of late antique/scholarly prose. "
            f"Translate ALL non-English text from {('detected source language' if not src or src=='auto' else src)} to {tgt}. "
            "Preserve ONLY tokens like [[G0]], [[G1]] exactly as-is (these are other scripts). "
            "Ignore running heads/footers and citation sigla if they slip through. "
            "Fix obvious OCR artifacts; keep punctuation and paragraphing. "
            "Return ONLY the translation text."
        )
        for i, t in enumerate(texts, 1):
            masked, mapping = _mask(t, self.mask_regex)
            params: Dict[str, Any] = {
                "model": self.model,
                "messages":[{"role":"system","content":sys_msg},
                            {"role":"user","content":masked}],
            }
            last_err: Optional[Exception] = None
            for attempt in range(6):
                try:
                    resp = self.client.chat.completions.create(**params)
                    text_out = (resp.choices[0].message.content or "").strip()
                    out.append(_unmask(text_out, mapping))
                    if self.verbose: print(f"  ✓ translated chunk {i}/{len(texts)}", flush=True)
                    break
                except Exception as e:
                    msg = str(e).lower(); last_err = e
                    if any(k in msg for k in ("rate","quota","429","temporar")):
                        time.sleep(2 ** attempt + random.random()); continue
                    raise
            else:
                raise RuntimeError(f"OpenAI translation failed after retries: {last_err}")
        return out

# ---------- writers ----------
try:
    from docx import Document  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.shared import Inches  # type: ignore
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # type: ignore
    _DOCX = True
except Exception:
    _DOCX = False

def _set_rtl(cell) -> None:
    try:
        tcPr = cell._tc.get_or_add_tcPr()
        bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), "1"); tcPr.append(bidi)
    except Exception:
        pass

def write_docx(out_path: Path, title: str, pages: List[Dict[str, Any]], rtl_original: bool, bilingual: bool) -> None:
    if not _DOCX: raise RuntimeError("pip install python-docx")
    doc = Document()
    h = doc.add_heading(title, 0); h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for page in pages:
        doc.add_heading(f"Page {page['page_number']}", level=1)
        if page.get("title") and len(page["title"]) <= 120:
            p = doc.add_paragraph(); run = p.add_run(page["title"]); run.italic = True
        if page.get("caput_norm"):
            doc.add_paragraph(page["caput_norm"])

        ncols = 2 if bilingual else 1
        table = doc.add_table(rows=1, cols=ncols)
        hdr = table.rows[0].cells
        hdr[0].text = "Original" if bilingual else "Translation"
        if bilingual: hdr[1].text = "Translation"
        try:
            if bilingual: hdr[0].width = Inches(4); hdr[1].width = Inches(4)
            else: hdr[0].width = Inches(8)
        except Exception: pass

        for (orig, trans) in page["pairs"]:
            r = table.add_row().cells
            if bilingual:
                r[0].text = orig; r[1].text = trans
                if rtl_original: _set_rtl(r[0])
            else:
                r[0].text = trans

        if page.get("foot_pairs"):
            doc.add_paragraph(); doc.add_heading("Footnotes", level=2)
            ft = doc.add_table(rows=1, cols=ncols)
            fh = ft.rows[0].cells
            fh[0].text = "Original" if bilingual else "Translation"
            if bilingual: fh[1].text = "Translation"
            for (fo, ft_) in page["foot_pairs"]:
                rr = ft.add_row().cells
                if bilingual:
                    rr[0].text = fo; rr[1].text = ft_
                    if rtl_original: _set_rtl(rr[0])
                else:
                    rr[0].text = ft_
        doc.add_page_break()
    doc.save(out_path)

def save_output(out_path: Path, title: str, pages: List[Dict[str, Any]],
                rtl_original: bool, out_format: str, bilingual: bool,
                pages_done: Optional[int] = None) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tmp = out_path.with_suffix(out_path.suffix + ".part")
    if out_format == "docx":
        write_docx(tmp, title, pages, rtl_original, bilingual)
    else:
        with tmp.open("w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n")
            for page in pages:
                f.write(f"## Page {page['page_number']}\n\n")
                if page.get("title"): f.write(f"*{page['title']}*\n\n")
                if page.get("caput_norm"): f.write(f"{page['caput_norm']}\n\n")
                for (orig, trans) in page["pairs"]:
                    if bilingual and orig.strip():
                        f.write("**Original**\n\n" + orig.strip() + "\n\n")
                    f.write("**Translation**\n\n" + trans.strip() + "\n\n---\n\n")
                if page.get("foot_pairs"):
                    f.write("### Footnotes\n\n")
                    for (fo, ft_) in page["foot_pairs"]:
                        if bilingual and fo.strip():
                            f.write("_Original footnote_\n\n" + fo.strip() + "\n\n")
                        f.write("_Translated footnote_\n\n" + ft_.strip() + "\n\n")
                    f.write("\n---\n\n")
    try:
        tmp.replace(out_path)
    except PermissionError:
        suffix = f".p{pages_done:04d}" if pages_done else ""
        fallback = out_path.with_suffix(out_path.suffix + f".checkpoint{suffix}")
        shutil.move(str(tmp), str(fallback))
        print(f"  ⚠ Could not overwrite '{out_path.name}'. Saved checkpoint: {fallback.name}", flush=True)

# ---------- main pipeline ----------
def process_pdf(
    pdf: Path,
    translator: OpenAITranslator,
    src: Optional[str],
    tgt: str,
    out_dir: Path,
    out_format: str,
    bilingual: bool,
    rtl_original: bool,
    autosave_every: int,
    resume_from: int,
    footnote_zone: float,
    footnote_min_lines: int,
    footnote_density: float,
    include_p_cue: bool,
    max_body_chars: int,
    max_footnote_chars: int,
    skip_ocr: bool,
    smart_ocr: bool,
    ocr_langs: str,
    ocr_psm: int,
    ocr_oversample: int,
    ocr_force_all: bool,
    extra_header_hints: List[str],
    title: Optional[str],
    verbose: bool,
    apparatus_mode: str,
    skip_front_matter: bool,
    normalize_caput: bool,
    stitch_pages: bool,
    footnote_mode: str,
) -> Path:
    title = title or pdf.name
    stem = safe_stem(pdf)
    out_path = out_dir / (f"{stem}_translated.docx" if out_format == "docx" else f"{stem}_translated.md")

    with tempfile.TemporaryDirectory() as td:
        text = ""
        if not skip_ocr:
            if verbose: print(f"Starting OCR for: {pdf.name}", flush=True)
            searchable, sidecar = ensure_searchable_pdf(
                pdf, ocr_langs, Path(td),
                force_all=ocr_force_all, oversample=ocr_oversample,
                psm=ocr_psm, verbose=verbose
            )
            if verbose: print("OCR complete.", flush=True)
            side_text = sidecar.read_text(encoding="utf-8", errors="ignore") if sidecar else ""
            if side_text.strip().startswith("[OCR skipped on page"):
                if verbose: print("  ↳ Sidecar reports 'OCR skipped' — falling back to text layer extraction.", flush=True)
                side_text = ""
            text = side_text or _extract_all_text(searchable)
        else:
            if verbose: print("Skipping OCR — using existing text layer.", flush=True)
            text = _extract_all_text(pdf)

        text = _strip_false_nonlatin_noise(text)

        if smart_ocr and not ocr_force_all:
            pages_probe = re.split(r"\f+", text) if "\f" in text else [text]
            needs = any(_needs_force_ocr(p) for p in pages_probe)
            if needs:
                if verbose: print("Smart-OCR: mentions Greek/Syriac but no real runs — forcing OCR.", flush=True)
                searchable2, sidecar2 = ensure_searchable_pdf(
                    pdf, ocr_langs, Path(td),
                    force_all=True, oversample=ocr_oversample, psm=ocr_psm, verbose=verbose
                )
                side_text2 = sidecar2.read_text(encoding="utf-8", errors="ignore") if sidecar2 else ""
                if side_text2.strip().startswith("[OCR skipped on page"): side_text2 = ""
                text = side_text2 or _extract_all_text(searchable2)
                text = _strip_false_nonlatin_noise(text)

    if not text.strip(): raise RuntimeError(f"No text extracted from {pdf}")
    pages_raw = re.split(r"\f+", text) if "\f" in text else [text]
    if verbose: print(f"Extracted ~{len(pages_raw)} page(s) of text.", flush=True)

    pages_out: List[Dict[str, Any]] = []
    translated_pages = 0
    carry = ""

    note_start_re = build_note_start_re(include_p_cue=include_p_cue)

    for page_number, raw_text in enumerate(pages_raw, start=1):
        t = raw_text
        for k, v in LIGS.items(): t = t.replace(k, v)
        t = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", t)
        t = re.sub(r"\s+([,;:.!?])", r"\1", t)
        t = _strip_marginal_numbers(t)

        cleaned, siphoned = _strip_editorial_and_headers(t, extra_hints=extra_header_hints, apparatus_mode=apparatus_mode)
        if skip_front_matter and _looks_front_matter(cleaned):
            if verbose: print(f"  (front-matter) Skipping page {page_number}.", flush=True)
            continue
        if page_number < resume_from:
            if verbose: print(f"(resume) Skipping page {page_number}/{len(pages_raw)}", flush=True)
            continue
        if not cleaned.strip():
            if verbose: print(f"(blank) Page {page_number} has no content.", flush=True)
            continue

        if verbose: print(f"\nTranslating page {page_number}/{len(pages_raw)} …", flush=True)

        heading, body_after_heading = peel_heading_block(cleaned)
        page_title = detect_title(cleaned) or heading
        caput_norm = _maybe_normalize_caput(page_title or "", normalize_caput)

        # if stitching, only carry over if next page's first nonempty line is NOT a note/app line
        if stitch_pages and carry:
            first_nonempty = next((x for x in body_after_heading.splitlines() if x.strip()), "")
            looks_note = bool(note_start_re.match(first_nonempty) or APPARATUS_LINE.match(first_nonempty))
            if not looks_note:
                body_after_heading = (carry + " " + body_after_heading).strip()
            carry = ""

        body_text, foot_text = split_body_and_footnotes(
            body_after_heading,
            zone_ratio=footnote_zone,
            min_lines=footnote_min_lines,
            density_req=footnote_density,
            note_start_re=note_start_re,
            mode=footnote_mode
        )
        if siphoned:
            foot_text = (foot_text + "\n" + siphoned).strip() if foot_text else siphoned

        if stitch_pages:
            tail = ""
            if not re.search(r"[.!?…»”)]\s*$", body_text):
                m_hy = re.search(r"([A-Za-z]{2,}\-\s*)$", body_text)
                if m_hy:
                    tail = m_hy.group(1)
                else:
                    m_tail = re.search(r"([^.?!…»”)]{10,180})$", body_text)
                    if m_tail:
                        tail = m_tail.group(1)
            if tail:
                carry = tail.strip("- \n\r\t")
                body_text = body_text[:-len(tail)].rstrip()

        body_paras = paragraphs_from_text(body_text)
        body_chunks = chunk_text(body_paras, max_chars=max_body_chars)
        if verbose: print(f"  body: {len(body_paras)} paragraph(s) → {len(body_chunks)} chunk(s)", flush=True)
        body_trs = translator.translate(body_chunks, src, tgt)

        pairs: List[Tuple[str, str]] = []
        for ch, tr in zip(body_chunks, body_trs):
            ps, ts = paragraphs_from_text(ch), paragraphs_from_text(tr)
            if len(ps) != len(ts): pairs.append(("\n\n".join(ps), "\n\n".join(ts)))
            else: pairs.extend(list(zip(ps, ts)))

        foot_pairs: Optional[List[Tuple[str, str]]] = None
        if foot_text:
            f_paras = paragraphs_from_text(foot_text)
            f_chunks = chunk_text(f_paras, max_chars=max_footnote_chars) if f_paras else []
            if f_chunks and verbose: print(f"  footnotes: {len(f_paras)} paragraph(s) → {len(f_chunks)} chunk(s)", flush=True)
            if f_chunks:
                f_trs = translator.translate(f_chunks, src, tgt)
                foot_pairs = []
                for ch, tr in zip(f_chunks, f_trs):
                    fps, fts = paragraphs_from_text(ch), paragraphs_from_text(tr)
                    if len(fps) != len(fts): foot_pairs.append(("\n\n".join(fps), "\n\n".join(fts)))
                    else: foot_pairs.extend(list(zip(fps, fts)))

        pages_out.append({
            "page_number": page_number,
            "title": page_title,
            "caput_norm": caput_norm,
            "pairs": pairs,
            "foot_pairs": foot_pairs
        })
        translated_pages += 1
        if translated_pages % max(1, autosave_every) == 0:
            print(f"  ↳ autosaving ({translated_pages} page(s) done)…", flush=True)
            save_output(out_path, title, pages_out, rtl_original, out_format, bilingual,
                        pages_done=resume_from - 1 + translated_pages)

    if verbose: print("\nFinal save…", flush=True)
    save_output(out_path, title, pages_out, rtl_original, out_format, bilingual,
                pages_done=resume_from - 1 + translated_pages)
    return out_path

# ---------- CLI ----------
def main() -> None:
    ap = argparse.ArgumentParser(description="Translate scholarly PDFs → bilingual DOCX/MD (clean, OCR-capable).")
    ap.add_argument("input", help="PDF file or directory")
    ap.add_argument("--out-dir", default="out")
    ap.add_argument("--src", default="auto", help="Source lang (e.g., auto, la, fr)")
    ap.add_argument("--tgt", default="en", help="Target lang (e.g., en, de)")
    ap.add_argument("--openai-model", default="gpt-5-mini")
    ap.add_argument("--format", choices=["docx","md"], default="docx")
    ap.add_argument("--translation-only", action="store_true")
    ap.add_argument("--rtl-original", action="store_true")
    ap.add_argument("--autosave-every", type=int, default=1)
    ap.add_argument("--resume-from", type=int, default=1)

    # conservative defaults
    ap.add_argument("--footnote-threshold", type=float, default=0.34,
                    help="Bottom-of-page window (0..1) to scan for footnote block.")
    ap.add_argument("--footnote-min-lines", type=int, default=4,
                    help="Require this many contiguous note-like lines to split.")
    ap.add_argument("--footnote-density", type=float, default=0.72,
                    help="Require ≥ this fraction of the bottom window to be note-like (0..1).")
    ap.add_argument("--footnote-split", choices=["strict","none"], default="strict",
                    help="Use 'none' to disable body↔footnote splitting entirely.")

    ap.add_argument("--note-include-p-cue", action="store_true",
                    help="Treat 'p. N' starts as note heads (default: OFF).")

    ap.add_argument("--max-body-chars", type=int, default=2200)
    ap.add_argument("--max-footnote-chars", type=int, default=1800)
    ap.add_argument("--extra-header-hints", default="CORPUS,SCRIPTORUM,CHRISTIANORUM,ORIENTALIUM,IMPRIMERIE,LOUVAIN,RÉIMPRESSION,ANASTATIQUE,SEVERI,GRAMMATICUM,PATRIARCHAE,ANTIOCHIAE,LIBER,SANCTI,ORATIO,PARS,POSTERIOR")
    ap.add_argument("--title", default=None)
    ap.add_argument("--skip-front-matter", action="store_true")
    ap.add_argument("--normalize-caput", action="store_true")

    ap.add_argument("--apparatus", choices=["drop","footnotes","keep"], default="footnotes",
                    help="How to handle editorial sigla like 'P. Gr., …'")

    ap.add_argument("--preserve-scripts", default="greek",
                    help="Comma list: greek, syriac, hebrew, coptic, armenian")

    g = ap.add_mutually_exclusive_group()
    g.add_argument("--skip-ocr", dest="skip_ocr", action="store_true", help="Use embedded text (default)")
    g.add_argument("--no-skip-ocr", dest="skip_ocr", action="store_false", help="Run OCR (ocrmypdf+tesseract)")
    ap.set_defaults(skip_ocr=True)
    ap.add_argument("--smart-ocr", dest="smart_ocr", action="store_true", help="Auto force-OCR if Greek/Syriac are missing (default).")
    ap.add_argument("--no-smart-ocr", dest="smart_ocr", action="store_false")
    ap.set_defaults(smart_ocr=True)
    ap.add_argument("--ocr-langs", default="eng,lat,fra,ell,grc", help="Tesseract langs (comma)")
    ap.add_argument("--ocr-psm", type=int, default=4)
    ap.add_argument("--ocr-oversample", type=int, default=600)
    ap.add_argument("--force-ocr-all", action="store_true")
    ap.add_argument("--stitch-pages", dest="stitch_pages", action="store_true", help="Stitch sentence tails across pages (default).")
    ap.add_argument("--no-stitch-pages", dest="stitch_pages", action="store_false")
    ap.set_defaults(stitch_pages=True)
    ap.add_argument("--verbose", action="store_true")

    args = ap.parse_args()

    inp = Path(args.input)
    if not inp.exists(): print(f"Not found: {inp}", file=sys.stderr); sys.exit(1)
    targets = [p for p in (inp.glob("*.pdf") if inp.is_dir() else [inp]) if p.suffix.lower()==".pdf"]
    if not targets: print("No PDFs found.", file=sys.stderr); sys.exit(1)

    scripts = [s.strip().lower() for s in args.preserve_scripts.split(",") if s.strip()]
    mask_rx = _compile_mask_regex(scripts)

    translator = OpenAITranslator(model=args.openai_model, verbose=args.verbose, mask_regex=mask_rx)
    out_dir = Path(args.out_dir); out_dir.mkdir(parents=True, exist_ok=True)
    bilingual = not args.translation_only
    extra_hints = [h.strip() for h in args.extra_header_hints.split(",") if h.strip()]

    print(f"Processing {len(targets)} PDF(s)…")
    for pdf in targets:
        try:
            out = process_pdf(
                pdf=pdf, translator=translator, src=args.src, tgt=args.tgt,
                out_dir=out_dir, out_format=args.format, bilingual=bilingual,
                rtl_original=args.rtl_original, autosave_every=args.autosave_every,
                resume_from=args.resume_from, footnote_zone=args.footnote_threshold,
                footnote_min_lines=args.footnote_min_lines, footnote_density=args.footnote_density,
                include_p_cue=args.note_include_p_cue,
                max_body_chars=args.max_body_chars, max_footnote_chars=args.max_footnote_chars,
                skip_ocr=args.skip_ocr, smart_ocr=args.smart_ocr, ocr_langs=args.ocr_langs,
                ocr_psm=args.ocr_psm, ocr_oversample=args.ocr_oversample,
                ocr_force_all=args.force_ocr_all, extra_header_hints=extra_hints,
                title=args.title, verbose=args.verbose, apparatus_mode=args.apparatus,
                skip_front_matter=args.skip_front_matter, normalize_caput=args.normalize_caput,
                stitch_pages=args.stitch_pages, footnote_mode=args.footnote_split
            )
            print(f"✔ {pdf.name} → {out}")
        except Exception as e:
            print(f"✖ {pdf.name} FAILED: {e}", file=sys.stderr)

if __name__ == "__main__":
    main()
